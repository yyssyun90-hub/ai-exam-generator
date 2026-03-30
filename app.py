import streamlit as st
import openai
from openai import OpenAI
import json
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib.patches import Ellipse
import math
import re
import fitz
from PIL import Image
import base64
from datetime import datetime

st.set_page_config(page_title="AI智能出卷助手", page_icon="📝", layout="wide")

# 初始化session_state
if "paper_json" not in st.session_state:
    st.session_state.paper_json = None
if "analyzed_style" not in st.session_state:
    st.session_state.analyzed_style = None
if "syllabus_units" not in st.session_state:
    st.session_state.syllabus_units = []
if "syllabus_content" not in st.session_state:
    st.session_state.syllabus_content = ""
if "today_calls" not in st.session_state:
    st.session_state.today_calls = 0
if "last_reset_date" not in st.session_state:
    st.session_state.last_reset_date = datetime.now().strftime("%Y-%m-%d")
if "selected_units" not in st.session_state:
    st.session_state.selected_units = []
if "extra_requirements" not in st.session_state:
    st.session_state.extra_requirements = ""
if "custom_prompt" not in st.session_state:
    # 默认 Prompt 模板（你可以自由修改）
    st.session_state.custom_prompt = """你是一位经验丰富的小学{subject}教师，请出一份{grade}{paper_type}试卷。

【科目特点】
- 华文：重点考察看拼音写字、组词、造句、阅读理解、看图写话
- 数学：重点考察计算题、应用题、图形题、填空题
- 健康教育：重点考察身体部位识别、食物分类、卫生习惯、安全知识

【出题范围】
只从以下内容出题：{unit_scope}

【难度】
{difficulty}

【额外要求】
{extra_requirements}

【试卷风格参考】
{style_str}

【图片处理规则】
1. 几何图形题 → 在题干前添加【自动绘图：图形描述】
2. 复杂场景题 → 在题干前添加【图X：详细场景描述】
3. 图片标记从【图1】开始递增

请生成一份完整试卷，保持与风格参考相似的格式。
题目要符合{grade}学生水平，语言生动有趣。

输出JSON格式：
{{
    "title": "试卷标题",
    "total_score": 总分,
    "sections": [
        {{
            "type": "选择题",
            "score_per_question": 每题分值,
            "questions": [
                {{
                    "number": 1,
                    "text": "题干",
                    "image_marker": "",
                    "image_type": "",
                    "image_description": "",
                    "options": ["A. xxx", "B. xxx", "C. xxx", "D. xxx"],
                    "answer": "A",
                    "explanation": "答案解析"
                }}
            ]
        }},
        {{
            "type": "填空题",
            "score_per_question": 每题分值,
            "questions": [
                {{
                    "number": 1,
                    "text": "题干",
                    "image_marker": "",
                    "image_type": "",
                    "image_description": "",
                    "answer": "答案",
                    "explanation": "答案解析"
                }}
            ]
        }},
        {{
            "type": "简答题/应用题",
            "score_per_question": 每题分值,
            "questions": [
                {{
                    "number": 1,
                    "text": "题干",
                    "image_marker": "",
                    "image_type": "",
                    "image_description": "",
                    "answer": "参考答案",
                    "explanation": "评分要点"
                }}
            ]
        }}
    ]
}}

只输出JSON，不要其他内容。"""

st.title("📝 AI 智能出卷助手")
st.caption("上传参考试卷和教学计划，AI学习风格后自动生成指定单元/课范围的试卷")


# ========== 辅助函数 ==========

def read_docx_content(file):
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text


def read_pdf_content(file):
    text = ""
    try:
        import pdfplumber
        file.seek(0)
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except:
        try:
            from PyPDF2 import PdfReader
            file.seek(0)
            pdf_reader = PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        except:
            pass
    return text


def track_api_call():
    today = datetime.now().strftime("%Y-%m-%d")
    if st.session_state.last_reset_date != today:
        st.session_state.last_reset_date = today
        st.session_state.today_calls = 0
    st.session_state.today_calls += 1


def get_remaining_calls():
    today = datetime.now().strftime("%Y-%m-%d")
    if st.session_state.last_reset_date != today:
        st.session_state.last_reset_date = today
        st.session_state.today_calls = 0
    return 50 - st.session_state.today_calls


def extract_units_from_syllabus(syllabus_content, client):
    """从教学计划中提取单元/课列表"""
    prompt = f"""
请从以下教学计划中提取所有的单元/章节/课的名称。

教学计划内容：
{syllabus_content[:3000]}

请识别内容的层级结构，可能是"第X单元"或"第X课"。
以JSON数组格式输出，例如：
["第1单元：认识数字", "第2单元：加减法", "第3课：10以内的加法", "第4课：10以内的减法"]

只输出JSON数组，不要其他内容。
"""
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "你是教学计划分析专家，只输出JSON数组。"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=500
        )
        units = json.loads(response.choices[0].message.content)
        return units if isinstance(units, list) else []
    except Exception as e:
        return ["第1单元", "第2单元", "第3单元", "第4单元", "第5单元", "第6单元"]


def analyze_paper_style(reference_files, client):
    """分析试卷风格"""
    papers_content = []
    for file in reference_files:
        if file.name.endswith('.pdf'):
            content = read_pdf_content(file)
        else:
            file.seek(0)
            content = read_docx_content(file)
        papers_content.append(content[:3000])
    
    combined = "\n---\n".join(papers_content)
    
    prompt = f"""
请分析以下参考试卷的风格特点，输出JSON格式。

参考试卷内容：
{combined}

请分析并输出：
{{
    "paper_structure": {{
        "total_score": 总分值,
        "section_types": ["题型1", "题型2"]
    }},
    "question_style": {{
        "language_style": "语言风格描述",
        "difficulty_distribution": "简单:中等:难的比例"
    }},
    "scoring_pattern": {{
        "choice_score": 每题分值,
        "fill_score": 每题分值,
        "essay_score": 每题分值
    }},
    "typical_topics": ["常见知识点1", "知识点2"]
}}

只输出JSON，不要其他内容。
"""
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "你是试卷分析专家，只输出JSON。"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=2000
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        return {"error": str(e)}


def generate_paper(style, grade, subject, paper_type, difficulty, extra_requirements, 
                   syllabus_content, selected_units, client, auto_draw=True, custom_prompt=""):
    """生成试卷 - 使用自定义 Prompt"""
    if not client:
        return {"error": "API Key未配置"}
    
    style_str = json.dumps(style, ensure_ascii=False)
    units_str = "、".join(selected_units)
    
    # 如果没有自定义 Prompt，使用默认的
    if not custom_prompt:
        custom_prompt = st.session_state.custom_prompt
    
    # 替换变量
    try:
        prompt = custom_prompt.format(
            subject=subject,
            grade=grade,
            paper_type=paper_type,
            unit_scope=units_str,
            difficulty=difficulty,
            extra_requirements=extra_requirements if extra_requirements else "无",
            style_str=style_str
        )
    except KeyError as e:
        return {"error": f"Prompt模板中的变量 {e} 不被支持，请使用: subject, grade, paper_type, unit_scope, difficulty, extra_requirements, style_str"}
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "你是小学试卷出题专家，只输出JSON。题目要符合小学水平，语言生动。"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=4000
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        return {"error": str(e)}


def draw_geometry(description):
    """根据描述绘制图形"""
    description = description.lower()
    
    fig, ax = plt.subplots(1, 1, figsize=(5, 4))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 8)
    ax.set_aspect('equal')
    ax.axis('off')
    
    if '长方形' in description or '矩形' in description:
        numbers = re.findall(r'\d+', description)
        length, width = (int(numbers[0]), int(numbers[1])) if len(numbers) >= 2 else (8, 4)
        rect = patches.Rectangle((1, 2), length, width, linewidth=2, edgecolor='black', facecolor='lightblue', alpha=0.7)
        ax.add_patch(rect)
        ax.set_xlim(0, length + 2)
        ax.set_ylim(0, width + 3)
        ax.annotate(f'{length}cm', xy=(1 + length/2, 1.5), ha='center')
        ax.annotate(f'{width}cm', xy=(0.5, 2 + width/2), va='center', rotation=90)
    
    elif '正方形' in description:
        numbers = re.findall(r'\d+', description)
        side = int(numbers[0]) if numbers else 5
        square = patches.Rectangle((2.5, 1.5), side, side, linewidth=2, edgecolor='black', facecolor='lightgreen', alpha=0.7)
        ax.add_patch(square)
        ax.annotate(f'{side}cm', xy=(2.5 + side/2, 1), ha='center')
    
    elif '三角形' in description:
        triangle = patches.Polygon([[3, 1], [7, 1], [5, 5]], linewidth=2, edgecolor='black', facecolor='lightcoral', alpha=0.7)
        ax.add_patch(triangle)
    
    elif '圆' in description or '圆形' in description:
        numbers = re.findall(r'\d+', description)
        radius = int(numbers[0]) / 2 if numbers else 2
        circle = patches.Circle((5, 4), radius, linewidth=2, edgecolor='black', facecolor='lightyellow', alpha=0.7)
        ax.add_patch(circle)
        ax.annotate(f'半径{radius}cm', xy=(5, 4 - radius - 0.3), ha='center')
    
    elif '云' in description:
        for pos in [(3,4,1.2), (4.5,4.5,1.5), (6,4,1.2), (4,3,1)]:
            circle = patches.Circle((pos[0], pos[1]), pos[2], edgecolor='black', facecolor='white', linewidth=1.5)
            ax.add_patch(circle)
    
    elif '太阳' in description:
        sun = patches.Circle((5, 5), 1.5, edgecolor='orange', facecolor='yellow', linewidth=2)
        ax.add_patch(sun)
        for angle in range(0, 360, 30):
            rad = math.radians(angle)
            ax.plot([5 + 1.5*math.cos(rad), 5 + 2.2*math.cos(rad)], 
                   [5 + 1.5*math.sin(rad), 5 + 2.2*math.sin(rad)], color='orange', linewidth=2)
    
    elif '树' in description:
        trunk = patches.Rectangle((4.5, 2), 1, 2, edgecolor='brown', facecolor='saddlebrown')
        crown = patches.Polygon([[3, 4], [7, 4], [5, 7]], edgecolor='green', facecolor='forestgreen', alpha=0.8)
        ax.add_patch(trunk)
        ax.add_patch(crown)
    
    elif '房子' in description:
        house = patches.Rectangle((3, 2), 4, 3, edgecolor='black', facecolor='lightyellow')
        roof = patches.Polygon([[2.5, 5], [7.5, 5], [5, 7]], edgecolor='brown', facecolor='lightcoral')
        door = patches.Rectangle((4.5, 2), 1, 1.5, edgecolor='black', facecolor='brown')
        window = patches.Circle((5.5, 3.5), 0.4, edgecolor='black', facecolor='lightblue')
        ax.add_patch(house)
        ax.add_patch(roof)
        ax.add_patch(door)
        ax.add_patch(window)
    
    elif '鱼' in description:
        body = Ellipse((5, 4), 3, 1.5, edgecolor='black', facecolor='orange', alpha=0.8)
        tail = patches.Polygon([[3.5, 4], [2, 3.5], [2, 4.5]], edgecolor='black', facecolor='orange')
        eye = patches.Circle((4, 4.2), 0.15, facecolor='black')
        ax.add_patch(body)
        ax.add_patch(tail)
        ax.add_patch(eye)
    
    elif '花' in description:
        for angle in range(0, 360, 45):
            rad = math.radians(angle)
            petal = patches.Circle((5 + 0.8*math.cos(rad), 4 + 0.8*math.sin(rad)), 0.4, 
                                   edgecolor='pink', facecolor='pink', alpha=0.7)
            ax.add_patch(petal)
        center = patches.Circle((5, 4), 0.3, edgecolor='brown', facecolor='yellow')
        stem = patches.Rectangle((4.8, 2), 0.4, 2, edgecolor='green', facecolor='green')
        ax.add_patch(center)
        ax.add_patch(stem)
    
    elif '苹果' in description:
        apple = patches.Circle((5, 4), 1.2, edgecolor='red', facecolor='red', alpha=0.7)
        leaf = patches.Polygon([[5.2, 5.2], [5.6, 5.5], [5.4, 5]], edgecolor='green', facecolor='green')
        stem = patches.Rectangle((4.9, 5), 0.2, 0.6, edgecolor='brown', facecolor='brown')
        ax.add_patch(apple)
        ax.add_patch(leaf)
        ax.add_patch(stem)
    
    elif '钟' in description or '时间' in description:
        clock = patches.Circle((5, 4), 2, edgecolor='black', facecolor='white', linewidth=2)
        ax.add_patch(clock)
        for hour in range(1, 13):
            angle = math.radians(90 - hour * 30)
            ax.plot([5 + 1.7*math.cos(angle), 5 + 1.9*math.cos(angle)],
                   [4 + 1.7*math.sin(angle), 4 + 1.9*math.sin(angle)], color='black', linewidth=1.5)
            ax.annotate(str(hour), xy=(5 + 1.5*math.cos(angle), 4 + 1.5*math.sin(angle)), 
                       ha='center', va='center', fontsize=10)
        numbers = re.findall(r'\d+', description)
        hour = int(numbers[0]) % 12 if numbers else 3
        minute = int(numbers[1]) if len(numbers) >= 2 else 0
        hour_angle = math.radians(90 - (hour * 30 + minute * 0.5))
        minute_angle = math.radians(90 - minute * 6)
        ax.plot([5, 5 + 1.0*math.cos(hour_angle)], [4, 4 + 1.0*math.sin(hour_angle)], 'k-', linewidth=3)
        ax.plot([5, 5 + 1.5*math.cos(minute_angle)], [4, 4 + 1.5*math.sin(minute_angle)], 'k-', linewidth=2)
    
    else:
        ax.text(5, 4, f'【需要绘制：{description}】', ha='center', va='center', fontsize=12, color='gray')
    
    buffer = BytesIO()
    plt.savefig(buffer, format='png', dpi=120, bbox_inches='tight', facecolor='white')
    buffer.seek(0)
    plt.close()
    return buffer


def create_word_document(paper_json, with_answers=False, auto_draw=True):
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    title = doc.add_heading(paper_json.get('title', '试卷'), level=0)
    title.alignment = 1
    
    info = doc.add_paragraph()
    info.add_run(f"总分：{paper_json.get('total_score', 100)}分")
    info.alignment = 1
    doc.add_paragraph()
    
    info_row = doc.add_paragraph()
    info_row.add_run("班级：_________________  姓名：_________________  学号：_________________")
    doc.add_paragraph()
    
    image_counter = 1
    
    for section in paper_json.get('sections', []):
        doc.add_heading(section.get('type', ''), level=2)
        doc.add_paragraph(f"（每题{section.get('score_per_question', 0)}分）")
        
        for q in section.get('questions', []):
            if q.get('image_marker') and auto_draw and q.get('image_type') == 'auto_draw':
                try:
                    img_buffer = draw_geometry(q.get('image_description', ''))
                    doc.add_picture(img_buffer, width=Inches(3))
                except:
                    doc.add_paragraph(f"【图{image_counter}：{q.get('image_description', '')}】")
                    image_counter += 1
            elif q.get('image_marker'):
                doc.add_paragraph(f"【图{image_counter}：{q.get('image_description', '')}】")
                image_counter += 1
            
            p = doc.add_paragraph()
            p.add_run(f"{q.get('number', '')}. {q.get('text', '')}")
            
            if 'options' in q and q['options']:
                for opt in q['options']:
                    doc.add_paragraph(f"   {opt}")
            
            if with_answers and 'answer' in q:
                ans = doc.add_paragraph()
                ans.add_run(f"【答案】{q['answer']}").italic = True
                ans.paragraph_format.left_indent = Pt(20)
            
            if not with_answers:
                doc.add_paragraph()
                doc.add_paragraph()
        
        doc.add_paragraph()
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ========== 主界面 UI ==========

# 侧边栏配置
with st.sidebar:
    st.header("🔐 配置")
    
    if "GITHUB_TOKEN" in st.secrets:
        api_key = st.secrets["GITHUB_TOKEN"]
        client = OpenAI(
            api_key=api_key,
            base_url="https://models.inference.ai.azure.com"
        )
        st.success("✅ GitHub Models API已配置")
        
        st.divider()
        st.header("📊 免费额度")
        
        remaining = get_remaining_calls()
        col1, col2 = st.columns(2)
        with col1:
            st.metric("今日已用", f"{st.session_state.today_calls} 次")
        with col2:
            st.metric("剩余", f"{remaining} 次")
        st.progress(min(st.session_state.today_calls / 50, 1.0))
        
        if remaining < 10:
            st.warning(f"⚠️ 剩余次数不足 {remaining} 次")
    else:
        st.warning("⚠️ 请在Streamlit Secrets中设置 GITHUB_TOKEN")
        client = None
        api_key = None
    
    st.divider()
    st.header("🎨 图片设置")
    auto_draw = st.checkbox("自动生成图形", value=True)
    st.info("复杂场景图片会生成【图X：描述】占位符")


# 三个标签页
tab1, tab2, tab3 = st.tabs(["📚 上传资料", "🎯 出卷设置", "📄 生成试卷"])

with tab1:
    st.subheader("📖 上传参考试卷")
    
    reference_papers = st.file_uploader(
        "上传之前的试卷（支持Word和PDF）",
        type=["docx", "pdf"],
        accept_multiple_files=True,
        help="上传2-3份，AI会学习题型、分值、语言风格"
    )
    
    st.subheader("📅 上传全年教学计划")
    syllabus_file = st.file_uploader(
        "上传教学计划（Word格式）",
        type=["docx"],
        help="AI会根据教学计划判断每个单元/课学什么"
    )
    
    if syllabus_file and st.button("📖 解析教学计划，提取单元/课列表"):
        if not client:
            st.error("请先配置 API Key")
        else:
            with st.spinner("正在解析教学计划..."):
                syllabus_content = read_docx_content(syllabus_file)
                st.session_state.syllabus_content = syllabus_content
                
                track_api_call()
                units = extract_units_from_syllabus(syllabus_content, client)
                st.session_state.syllabus_units = units
                
                st.success(f"成功提取 {len(units)} 个单元/课")
                with st.expander("查看列表"):
                    for i, unit in enumerate(units, 1):
                        st.write(f"{i}. {unit}")
    
    if reference_papers:
        st.success(f"已上传 {len(reference_papers)} 份参考试卷")

with tab2:
    st.subheader("🎯 出卷参数")
    
    col1, col2 = st.columns(2)
    with col1:
        grade_level = st.selectbox(
            "年级",
            ["一年级", "二年级", "三年级", "四年级", "五年级", "六年级"],
            key="grade_level"
        )
        subject = st.selectbox("科目", ["华文", "数学", "健康教育"], key="subject")
    
    with col2:
        paper_type = st.selectbox("试卷类型", ["单元测验", "期中考试", "期末考试", "模拟练习"], key="paper_type")
        difficulty = st.select_slider("难度", ["简单", "中等偏易", "中等", "中等偏难", "难"], value="中等", key="difficulty")
    
    # ========== 出题范围选择 ==========
    st.subheader("📚 出题范围")
    
    if st.session_state.syllabus_units:
        st.write("请勾选要出题的单元或课：")
        
        selected_units = []
        for item in st.session_state.syllabus_units:
            key = f"scope_{item}"
            if key not in st.session_state:
                st.session_state[key] = True  # 默认全选
            if st.checkbox(item, key=key):
                selected_units.append(item)
        
        st.session_state.selected_units = selected_units
        
        if selected_units:
            st.success(f"已选择 {len(selected_units)} 个内容")
        else:
            st.warning("请至少选择一个单元或课")
    else:
        st.info("请先在「上传资料」标签页上传教学计划并点击「解析教学计划」")
        st.session_state.selected_units = []
    
    # ========== 自定义 Prompt 模板 ==========
    st.subheader("✏️ 自定义出题指令")
    st.caption("你可以修改下面的模板，AI会按照你的要求出题。{科目}、{年级}等变量会被自动替换。")
    
    # 直接绑定 session_state
    custom_prompt = st.text_area(
        "出题指令模板",
        value=st.session_state.custom_prompt,
        height=300,
        help="可用变量：{subject} {grade} {paper_type} {unit_scope} {difficulty} {extra_requirements} {style_str}",
        key="custom_prompt_textarea"
    )
    
    # 实时保存
    if custom_prompt != st.session_state.custom_prompt:
        st.session_state.custom_prompt = custom_prompt
    
    col_save, col_reset = st.columns(2)
    with col_save:
        if st.button("💾 保存当前模板"):
            st.success("模板已保存！")
    with col_reset:
        if st.button("🔄 恢复默认模板"):
            st.session_state.custom_prompt = """你是一位经验丰富的小学{subject}教师，请出一份{grade}{paper_type}试卷。

【科目特点】
- 华文：重点考察看拼音写字、组词、造句、阅读理解、看图写话
- 数学：重点考察计算题、应用题、图形题、填空题
- 健康教育：重点考察身体部位识别、食物分类、卫生习惯、安全知识

【出题范围】
只从以下内容出题：{unit_scope}

【难度】
{difficulty}

【额外要求】
{extra_requirements}

【试卷风格参考】
{style_str}

请输出JSON格式，包含选择题、填空题、简答题/应用题。
题目要符合{grade}学生水平，语言生动有趣。"""
            st.rerun()
    
    # ========== 额外要求 ==========
    st.subheader("📋 额外要求")
    extra_requirements = st.text_area(
        "这次出卷的特殊要求",
        placeholder="例如：多出应用题、重点考察计算能力、多出看图写话题...",
        height=80,
        key="extra_requirements_input"
    )
    st.session_state.extra_requirements = extra_requirements
    
    # ========== 分析风格按钮 ==========
    if reference_papers and st.button("🔍 分析试卷风格", type="secondary"):
        if not client:
            st.error("请先配置 API Key")
        elif get_remaining_calls() <= 0:
            st.error("❌ 今日免费额度已用完，请明天再试")
        else:
            with st.spinner("AI正在分析试卷风格..."):
                track_api_call()
                style_analysis = analyze_paper_style(reference_papers, client)
                st.session_state.analyzed_style = style_analysis
                st.success("风格分析完成！")
                with st.expander("查看分析结果"):
                    st.json(style_analysis)

with tab3:
    # 确保必要变量存在
    if "selected_units" not in st.session_state:
        st.session_state.selected_units = []
    if "analyzed_style" not in st.session_state:
        st.session_state.analyzed_style = None
    
    if not st.session_state.analyzed_style:
        st.info("请先在「上传资料」标签页上传参考试卷，并点击「分析试卷风格」")
    elif not st.session_state.selected_units:
        st.info("请先在「出卷设置」标签页选择出题范围（单元或课）")
    else:
        # 获取最新的参数
        grade_level = st.session_state.get("grade_level", "三年级")
        subject = st.session_state.get("subject", "数学")
        paper_type = st.session_state.get("paper_type", "单元测验")
        difficulty = st.session_state.get("difficulty", "中等")
        extra_requirements = st.session_state.get("extra_requirements", "")
        
        st.subheader("🚀 生成新试卷")
        st.write(f"**出题范围：** {', '.join(st.session_state.selected_units)}")
        st.write(f"**学习风格来源：** {len(reference_papers) if reference_papers else 0} 份参考试卷")
        
        if st.button("📝 生成试卷", type="primary", use_container_width=True):
            if get_remaining_calls() <= 0:
                st.error("❌ 今日免费额度已用完，请明天再试")
            elif not reference_papers:
                st.error("请先上传参考试卷")
            else:
                with st.spinner("AI正在生成试卷..."):
                    track_api_call()
                    paper = generate_paper(
                        style=st.session_state.analyzed_style,
                        grade=grade_level,
                        subject=subject,
                        paper_type=paper_type,
                        difficulty=difficulty,
                        extra_requirements=extra_requirements,
                        syllabus_content=st.session_state.syllabus_content,
                        selected_units=st.session_state.selected_units,
                        client=client,
                        auto_draw=auto_draw,
                        custom_prompt=st.session_state.custom_prompt
                    )
                    st.session_state.paper_json = paper
                    st.success("试卷生成成功！")
    
    if st.session_state.paper_json:
        st.divider()
        st.subheader("📄 试卷预览")
        
        paper = st.session_state.paper_json
        
        if "error" in paper:
            st.error(f"生成失败：{paper['error']}")
        else:
            st.markdown(f"### {paper.get('title', '试卷')}")
            st.markdown(f"**总分：{paper.get('total_score', 100)}分**")
            st.markdown("---")
            
            for section in paper.get('sections', []):
                st.markdown(f"#### {section.get('type', '')}（每题{section.get('score_per_question', 0)}分）")
                for q in section.get('questions', []):
                    st.markdown(f"**{q.get('number', '')}.** {q.get('text', '')}")
                    if q.get('image_marker'):
                        st.info(f"🖼️ {q['image_marker']}")
                    if 'options' in q and q['options']:
                        for opt in q['options']:
                            st.markdown(f"&nbsp;&nbsp;&nbsp;{opt}")
                    st.markdown("")
                st.markdown("---")
            
            # 获取最新的年级、科目用于文件名
            grade_level = st.session_state.get("grade_level", "三年级")
            subject = st.session_state.get("subject", "数学")
            paper_type = st.session_state.get("paper_type", "单元测验")
            
            col_dl1, col_dl2 = st.columns(2)
            with col_dl1:
                if st.button("📥 下载学生卷"):
                    doc_buffer = create_word_document(
                        st.session_state.paper_json, 
                        with_answers=False,
                        auto_draw=auto_draw
                    )
                    st.download_button(
                        label="点击下载",
                        data=doc_buffer,
                        file_name=f"{grade_level}{subject}{paper_type}_学生卷.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_student"
                    )
            with col_dl2:
                if st.button("📥 下载教师卷（含答案）"):
                    doc_buffer = create_word_document(
                        st.session_state.paper_json, 
                        with_answers=True,
                        auto_draw=auto_draw
                    )
                    st.download_button(
                        label="点击下载",
                        data=doc_buffer,
                        file_name=f"{grade_level}{subject}{paper_type}_教师卷.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_teacher"
                    )
